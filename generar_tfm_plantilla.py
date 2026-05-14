"""
Genera el TFM reestructurado según la plantilla UNIR (plantilla_individual.docx).
Reorganiza el contenido del borrador actual y aplica las directrices del tutor:
  - Anglicismos minimizados (análisis de sentimiento, bancos de pruebas, etc.)
  - Acrónimos definidos en primera mención
  - Embeddings justificado como anglicismo necesario
  - Contexto de alucinaciones LLM en motivación del robo-advisor
  - Dataset de evaluación como contribución del trabajo
  - Estructura: plantilla_individual (Tipo 2 - Desarrollo de software)

Ejecutar: python generar_tfm_plantilla.py
Salida:  TFM_alertas_inversion_reestructurado.docx
"""

import copy
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ──────────────────────────────────────────────────────────────────
# 1. Cargar la plantilla y el documento fuente
# ──────────────────────────────────────────────────────────────────

PLANTILLA = r"plantilla_individual.docx"
FUENTE    = r"TFM_alertas_inversion_estado_de_la_cuestion.docx"
SALIDA    = r"TFM_alertas_inversion_reestructurado.docx"

plantilla = Document(PLANTILLA)
fuente    = Document(FUENTE)


# ──────────────────────────────────────────────────────────────────
# 2. Extraer contenido del documento fuente, agrupado por sección
# ──────────────────────────────────────────────────────────────────

def extraer_secciones(doc):
    """Devuelve dict {nombre_seccion: [texto_parrafo, ...]}"""
    secs = {}
    cur = "__PREAMBULO__"
    secs[cur] = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        style = p.style.name if p.style else ""
        if "Heading 1" in style and txt:
            cur = txt
            secs[cur] = []
        elif "Heading 2" in style and txt:
            cur = txt
            secs[cur] = []
        elif txt:
            secs[cur].append(txt)
    return secs

def extraer_tablas(doc):
    """Devuelve lista de tablas como listas de listas de strings."""
    tablas = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text for c in r.cells])
        tablas.append(rows)
    return tablas

secciones = extraer_secciones(fuente)
tablas_orig = extraer_tablas(fuente)


# ──────────────────────────────────────────────────────────────────
# 3. Funciones auxiliares para el nuevo documento
# ──────────────────────────────────────────────────────────────────

def sec(nombre):
    """Obtener párrafos de una sección del doc fuente."""
    return secciones.get(nombre, [])

def juntar(*nombres):
    """Concatenar párrafos de varias secciones."""
    resultado = []
    for n in nombres:
        resultado.extend(sec(n))
    return resultado


# ──────────────────────────────────────────────────────────────────
# 4. Sustituciones lingüísticas (directrices del tutor)
# ──────────────────────────────────────────────────────────────────

# Acrónimos: se definen la primera vez, luego se usa solo la sigla.
# Usamos un set para rastrear qué acrónimos ya se han definido.
acronimos_definidos = set()

ACRONIMOS = {
    "NLP":  ("procesamiento de lenguaje natural", "NLP", "inglés"),
    "LLM":  ("modelo de lenguaje de gran escala", "LLM", "inglés"),
    "LLMs": ("modelos de lenguaje de gran escala", "LLM", "inglés"),
    "NLI":  ("inferencia de lenguaje natural", "NLI", "inglés"),
    "NER":  ("reconocimiento de entidades nombradas", "NER", "inglés"),
    "API":  ("interfaz de programación de aplicaciones", "API", "inglés"),
    "APIs": ("interfaces de programación de aplicaciones", "API", "inglés"),
    "RSS":  ("sindicación realmente simple", "RSS", "inglés"),
    "ISIN": ("código internacional de identificación de valores", "ISIN", "inglés"),
    "SEC":  ("Comisión de Bolsa y Valores de Estados Unidos", "SEC", "inglés"),
    "CNMV": ("Comisión Nacional del Mercado de Valores", "CNMV", "español"),
    "TFM":  ("Trabajo Fin de Máster", "TFM", "español"),
}

# Anglicismos → equivalentes en español
ANGLICISMOS = [
    # (patrón regex, reemplazo) — se aplican en orden
    (r'\bsentiment analysis\b', 'análisis de sentimiento'),
    (r'\bSentiment Analysis\b', 'Análisis de Sentimiento'),
    (r'\bsentiment\b(?! analysis)', 'sentimiento'),
    (r'\bportfolio optimization\b', 'optimización de carteras'),
    (r'\bPortfolio Optimization\b', 'Optimización de Carteras'),
    (r'\bbenchmarks\b', 'bancos de pruebas'),
    (r'\bbenchmark\b', 'banco de pruebas'),
    (r'\bforecasting\b', 'predicción'),
    (r'\bForecasting\b', 'Predicción'),
    (r'\brisk management\b', 'gestión de riesgos'),
    (r'\bRisk Management\b', 'Gestión de Riesgos'),
    (r'\bportfolio management\b', 'gestión de carteras'),
    (r'\bPortfolio Management\b', 'Gestión de Carteras'),
    (r'\bcompliance\b', 'cumplimiento normativo'),
    (r'\bCompliance\b', 'Cumplimiento normativo'),
    (r'\btrading\b', 'negociación'),
    (r'\bscreener(?:s)?\b', 'filtro(s) de selección'),
    (r'\bfine-tuning\b', 'ajuste fino'),
    (r'\bFine-tuning\b', 'Ajuste fino'),
    (r'\bdataset(?:s)?\b', 'conjunto(s) de datos'),
    (r'\bDataset(?:s)?\b', 'Conjunto(s) de datos'),
    (r'\bpipeline\b', 'cadena de procesamiento'),
    (r'\bPipeline\b', 'Cadena de procesamiento'),
    (r'\btoken(?:s)?\b', 'token(es)'),
    (r'\bscoring\b', 'puntuación'),
    (r'\bbacktest(?:ing)?\b', 'prueba retrospectiva'),
    (r'\bframework\b', 'marco'),
    (r'\bFramework\b', 'Marco'),
    (r'\bopen source\b', 'código abierto'),
    (r'\bOpen Source\b', 'Código Abierto'),
    (r'\bfeedback\b', 'retroalimentación'),
    (r'\bFeedback\b', 'Retroalimentación'),
    (r'\bweb scraping\b', 'extracción web'),
    (r'\bspam\b', 'correo no deseado'),
    (r'\bsmart\b(?!\s*\()', 'inteligente'),
    (r'\bestado del arte\b', 'estado de la cuestión'),
    (r'\bEstado del arte\b', 'Estado de la cuestión'),
]

# Términos que se MANTIENEN en inglés pero se justifican en primera mención
ANGLICISMOS_JUSTIFICADOS = {
    "embeddings": (
        'representaciones vectoriales densas (embeddings, término adoptado de forma '
        'generalizada en la comunidad investigadora al no existir un equivalente '
        'consolidado en español; véase Mikolov et al., 2013; Reimers y Gurevych, 2019)'
    ),
    "robo-advisor": None,  # aceptado, no requiere justificación
    "robo-advisors": None,
    "software": None,
    "online": None,
    "F1": None,
    "recall": None,
    "precision": None,
    "FinBERT": None,
    "BERT": None,
    "ticker": None,
    "tickers": None,
    "hash": None,
    "score": None,
    "Deep Learning": None,
}

anglicismos_justificados_usados = set()


def aplicar_sustituciones(texto):
    """Aplica todas las sustituciones lingüísticas a un párrafo."""
    global acronimos_definidos, anglicismos_justificados_usados

    # 1. Acrónimos: definir en primera aparición
    for sigla, (expansion, sigla_base, idioma) in ACRONIMOS.items():
        pattern = r'\b' + re.escape(sigla) + r'\b'
        if re.search(pattern, texto) and sigla_base not in acronimos_definidos:
            if idioma == "inglés":
                repl = f"{expansion} ({sigla}, por sus siglas en inglés)"
            else:
                repl = f"{expansion} ({sigla})"
            # Solo reemplazar la primera ocurrencia
            texto = re.sub(pattern, repl, texto, count=1)
            acronimos_definidos.add(sigla_base)

    # 2. Anglicismos justificados: expandir solo primera vez
    for term, justificacion in ANGLICISMOS_JUSTIFICADOS.items():
        if justificacion and term in texto and term not in anglicismos_justificados_usados:
            texto = texto.replace(term, justificacion, 1)
            anglicismos_justificados_usados.add(term)

    # 3. Anglicismos → español
    for pattern, repl in ANGLICISMOS:
        texto = re.sub(pattern, repl, texto, flags=re.IGNORECASE)

    return texto


# ──────────────────────────────────────────────────────────────────
# 5. Construir el documento nuevo basado en la plantilla
# ──────────────────────────────────────────────────────────────────

# Usamos la plantilla como base: copiamos el documento completo para
# conservar estilos, fuentes, formato de página, encabezados, etc.
# Luego vaciamos todo el body y lo rellenamos con el contenido nuevo.

doc = Document(PLANTILLA)

# Limpiar todo el contenido existente de la plantilla excepto la primera
# sección de portada (que contiene los datos de la universidad)
# Estrategia: eliminar todos los párrafos y reconstruir, preservando
# las propiedades de sección (sectPr) que definen márgenes/tamaño de página.

body = doc.element.body
# Guardar sectPr antes de limpiar
from lxml import etree
sect_prs = body.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr')
saved_sect_pr = [copy.deepcopy(sp) for sp in sect_prs]

# Eliminar todos los hijos del body
for child in list(body):
    body.remove(child)

# Restaurar sectPr
for sp in saved_sect_pr:
    body.append(sp)


def add_paragraph(text, style="Normal"):
    """Añadir un párrafo con estilo y sustituciones lingüísticas."""
    cleaned = aplicar_sustituciones(text)
    try:
        p = doc.add_paragraph(cleaned, style=style)
    except KeyError:
        p = doc.add_paragraph(cleaned, style="Normal")
    return p


def add_heading(text, level=1):
    """Añadir un encabezado."""
    return doc.add_heading(text, level=level)


def add_list_item(text):
    """Añadir un elemento de lista."""
    return add_paragraph(text, style="List Paragraph")


def add_blank():
    """Añadir párrafo vacío."""
    return doc.add_paragraph("")


def add_table(headers, rows):
    """Añadir una tabla simple."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Encabezados
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    # Filas
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = aplicar_sustituciones(val)
    return table


# ──────────────────────────────────────────────────────────────────
# 6. PORTADA
# ──────────────────────────────────────────────────────────────────

p = doc.add_paragraph("Universidad Internacional de La Rioja")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Escuela Superior de Ingeniería y Tecnología")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_blank()
p = doc.add_paragraph("Máster Universitario en Inteligencia Artificial")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_blank()
add_blank()

titulo = (
    "Sistema de alertas inteligentes por noticias para carteras de inversión "
    "mediante detección de eventos y estimación de impacto con procesamiento "
    "de lenguaje natural"
)
p = doc.add_paragraph(titulo)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.bold = True
    run.font.size = Pt(16)

add_blank()
p = doc.add_paragraph("Trabajo Fin de Estudios")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_blank()
p = doc.add_paragraph("Presentado por: Rubén Querol Cervantes")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Director/a: [Nombre del director/a]")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Ciudad: [Ciudad]")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Fecha: [Fecha de presentación]")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────
# 7. RESUMEN (150-300 palabras, español)
# ──────────────────────────────────────────────────────────────────

add_heading("Resumen", level=0)

resumen_es = (
    "La inversión financiera contemporánea se enfrenta a una sobreabundancia "
    "informativa que dificulta al inversor individual la discriminación de "
    "noticias realmente relevantes para sus posiciones. Este Trabajo Fin de "
    "Máster (TFM) presenta el diseño, implementación y evaluación de "
    "InvestAIlert, un sistema inteligente de alertas por noticias orientado "
    "a carteras de inversión declaradas por el usuario. El sistema integra "
    "técnicas de procesamiento de lenguaje natural (NLP, por sus siglas en "
    "inglés) financiero —análisis de sentimiento con FinBERT, reconocimiento "
    "de entidades, representaciones vectoriales densas (embeddings) para "
    "deduplicación semántica y clasificación de eventos mediante inferencia "
    "de lenguaje natural (NLI) con BART-MNLI— con una lógica de "
    "personalización que evalúa la relevancia de cada noticia en función de "
    "la composición concreta de la cartera."
)
resumen_es2 = (
    "La metodología combina un desarrollo incremental de software con una "
    "evaluación cuantitativa basada en un corpus de 40 noticias etiquetadas "
    "manualmente, tres carteras de referencia y un estudio de ablación que "
    "compara cuatro variantes del sistema (reglas, híbrido, híbrido con NLI "
    "y completo con modelos de lenguaje de gran escala). Los resultados de "
    "la variante basada en reglas muestran un F1 de 0,935 en detección de "
    "relevancia (precisión perfecta) y un F1 de 1,0 en identificación de "
    "activos afectados. El sistema genera alertas explicables con "
    "trazabilidad de fuente, estimación de dirección e impacto, y control "
    "de ruido mediante deduplicación semántica."
)
resumen_es3 = (
    "La principal contribución reside en la integración de módulos de NLP "
    "financiero dentro de un marco de personalización por cartera, un "
    "enfoque que la literatura previa no ha abordado de forma conjunta. "
    "Adicionalmente, se aporta un conjunto de datos etiquetado bilingüe "
    "para la evaluación de sistemas de alertas financieras personalizadas."
)

# Reset acrónimos para el Abstract (el Resumen ya los define en su propia burbuja)
add_paragraph(resumen_es)
add_paragraph(resumen_es2)
add_paragraph(resumen_es3)

add_blank()
add_paragraph(
    "Palabras clave: alertas financieras, procesamiento de lenguaje natural, "
    "análisis de sentimiento, detección de eventos, carteras de inversión"
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────
# 8. ABSTRACT (150-300 words, English)
# ──────────────────────────────────────────────────────────────────

add_heading("Abstract", level=0)

abstract_en = (
    "Contemporary financial investing is characterized by an unprecedented "
    "abundance of textual information, making it difficult for individual "
    "investors to identify news that is genuinely material to their "
    "holdings. This Master's Thesis presents the design, implementation, "
    "and evaluation of InvestAIlert, an intelligent news-driven alert "
    "system tailored to user-declared investment portfolios. The system "
    "integrates financial natural language processing (NLP) techniques "
    "— sentiment analysis with FinBERT, named entity recognition, dense "
    "vector representations (embeddings) for semantic deduplication, and "
    "event classification via natural language inference (NLI) with "
    "BART-MNLI — with a personalization layer that assesses the relevance "
    "of each news item against the specific composition of the portfolio."
)
abstract_en2 = (
    "The methodology combines incremental software development with "
    "quantitative evaluation based on a manually labeled corpus of 40 news "
    "items, three reference portfolios, and an ablation study comparing "
    "four system variants (rules, hybrid, hybrid with NLI, and full with "
    "large language models). Results for the rule-based variant show an F1 "
    "score of 0.935 for relevance detection (perfect precision) and an F1 "
    "of 1.0 for affected asset identification. The system generates "
    "explainable alerts with source traceability, direction and impact "
    "estimation, and noise control via semantic deduplication."
)
abstract_en3 = (
    "The main contribution lies in integrating financial NLP modules within "
    "a portfolio-aware personalization framework, an approach not "
    "previously addressed holistically in the literature. Additionally, "
    "a bilingual labeled dataset for evaluating personalized financial "
    "alert systems is provided as a standalone contribution."
)

add_paragraph(abstract_en)
add_paragraph(abstract_en2)
add_paragraph(abstract_en3)

add_blank()
add_paragraph(
    "Keywords: financial alerts, natural language processing, sentiment "
    "analysis, event detection, investment portfolios"
)

doc.add_page_break()

# ──────────────────────────────────────────────────────────────────
# NOTA: El índice de contenidos, figuras y tablas se generan
# automáticamente en Word. Dejamos marcadores de posición.
# ──────────────────────────────────────────────────────────────────

add_heading("Índice de contenidos", level=0)
add_paragraph("[El índice de contenidos se generará automáticamente en Word: "
              "Referencias → Tabla de contenido → Insertar tabla de contenido]")
doc.add_page_break()

add_heading("Índice de figuras", level=0)
add_paragraph("[El índice de figuras se generará automáticamente en Word: "
              "Referencias → Insertar tabla de ilustraciones → Categoría: Figura]")

add_heading("Índice de tablas", level=0)
add_paragraph("[El índice de tablas se generará automáticamente en Word: "
              "Referencias → Insertar tabla de ilustraciones → Categoría: Tabla]")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  CAPÍTULO 1: INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════

# Reset de acrónimos para el cuerpo principal del texto
acronimos_definidos = set()
anglicismos_justificados_usados = set()

add_heading("Introducción", level=1)

# Párrafo introductorio (del original §1)
for p_text in sec("1. Introducción"):
    add_paragraph(p_text)

# ── 1.1 Motivación ──
add_heading("Motivación", level=2)

# Del original §2 (Problema abordado y motivación)
for p_text in sec("2. Problema abordado y motivación"):
    add_paragraph(p_text)

# NUEVO: Párrafo sobre alucinaciones de LLM como motivación
add_paragraph(
    "Un aspecto adicional que refuerza la motivación de este trabajo es la "
    "necesidad de contrarrestar las limitaciones de los modelos de lenguaje "
    "de gran escala cuando se emplean en contextos financieros. La "
    "literatura reciente ha documentado de forma creciente el fenómeno de "
    "las alucinaciones en estos modelos: la generación de información "
    "plausible pero no sustentada en evidencia real (Ji et al., 2023). En "
    "el ámbito financiero, donde las decisiones se apoyan en datos "
    "verificables y fuentes primarias, este riesgo resulta particularmente "
    "problemático. Un sistema de alertas basado en fuentes trazables y en "
    "técnicas de extracción de información —en lugar de en la generación "
    "libre de texto— constituye una respuesta deliberada a esta limitación: "
    "la información que recibe el inversor procede de noticias reales, "
    "procesadas y verificadas, no de contenido fabricado por un modelo "
    "generativo. Esta orientación hacia la información basada en evidencias "
    "es uno de los principios rectores del diseño del sistema propuesto."
)

# ── 1.2 Planteamiento del trabajo ──
add_heading("Planteamiento del trabajo", level=2)

# Síntesis del hueco detectado + propuesta (de §7 y §8)
add_paragraph(
    "La revisión bibliográfica que se desarrolla en el capítulo 2 permite "
    "identificar un hueco de investigación claro: la literatura dispone de "
    "trabajos maduros sobre robo-advisors, sobre análisis de sentimiento "
    "financiero, sobre extracción de eventos y sobre agentes basados en "
    "modelos de lenguaje de gran escala, pero estos avances no han convergido "
    "en un sistema integrado que condicione la relevancia de una noticia a "
    "la composición de una cartera concreta."
)

# Del original §7 (Hueco de investigación) - condensado
hueco_parrafos = sec("7. Hueco de investigación detectado")
if len(hueco_parrafos) >= 2:
    add_paragraph(hueco_parrafos[0])

# Del original §8 (Propuesta del sistema) - condensado
propuesta_parrafos = sec("8. Propuesta del sistema")
for p_text in propuesta_parrafos[:3]:
    add_paragraph(p_text)

# ── 1.3 Estructura del trabajo ──
add_heading("Estructura del trabajo", level=2)

add_paragraph(
    "El presente documento se estructura en cinco capítulos, además de las "
    "referencias bibliográficas y un anexo. A continuación se describe "
    "brevemente el contenido de cada uno."
)
add_paragraph(
    "El capítulo 1 (Introducción) ha presentado la motivación del trabajo, "
    "el problema detectado y el planteamiento general de la solución "
    "propuesta."
)
add_paragraph(
    "El capítulo 2 (Contexto y estado de la cuestión) revisa en profundidad "
    "la literatura sobre robo-advisors, adopción y confianza del usuario, "
    "agentes financieros basados en modelos de lenguaje de gran escala, "
    "procesamiento de lenguaje natural financiero aplicado a noticias y "
    "eventos, y optimización de carteras como marco complementario. Se "
    "concluye con el análisis comparativo y la identificación del hueco de "
    "investigación."
)
add_paragraph(
    "El capítulo 3 (Objetivos concretos y metodología de trabajo) formula "
    "el objetivo general y los objetivos específicos del trabajo siguiendo "
    "la metodología SMART, y describe la metodología de desarrollo "
    "incremental seguida."
)
add_paragraph(
    "El capítulo 4 (Desarrollo específico de la contribución) detalla la "
    "identificación de requisitos, la arquitectura y los módulos del "
    "sistema implementado, y presenta la evaluación cuantitativa con el "
    "estudio de ablación y el corpus etiquetado como contribución propia."
)
add_paragraph(
    "El capítulo 5 (Conclusiones y trabajo futuro) resume las aportaciones "
    "del trabajo, discute el grado de consecución de los objetivos y señala "
    "las líneas de trabajo futuro más prometedoras."
)
add_paragraph(
    "Finalmente, el Anexo A proporciona la referencia al repositorio de "
    "código fuente y a los datos empleados en la evaluación."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  CAPÍTULO 2: CONTEXTO Y ESTADO DE LA CUESTIÓN
# ══════════════════════════════════════════════════════════════════

add_heading("Contexto y estado de la cuestión", level=1)

add_paragraph(
    "Este capítulo revisa el estado actual de la investigación en las tres "
    "líneas que convergen en el sistema propuesto: el asesoramiento "
    "financiero automatizado (robo-advisors), el procesamiento de lenguaje "
    "natural aplicado a finanzas, y los agentes basados en modelos de "
    "lenguaje de gran escala. Se analiza también la optimización de "
    "carteras como marco complementario y se concluye con la identificación "
    "del hueco de investigación que justifica la presente contribución."
)

# ── 2.1 Contexto del problema ──
add_heading("Contexto del problema", level=2)

# Robo-advisors tradicionales (§3.1)
add_paragraph(
    "2.1.1. Robo-advisors tradicionales: asignación, perfilado y "
    "automatización del consejo"
)
for p_text in sec("3.1. Robo-advisors tradicionales: foco en asignación, "
                   "perfilado y automatización del consejo"):
    add_paragraph(p_text)

# Adopción y confianza (§3.2)
add_blank()
add_paragraph(
    "2.1.2. Adopción, confianza y límites de personalización"
)
for p_text in sec("3.2. Adopción, confianza y límites de personalización"):
    add_paragraph(p_text)

# Optimización de carteras (§5)
add_blank()
add_paragraph(
    "2.1.3. Optimización de carteras como marco complementario"
)
for p_text in sec("5. Optimización de carteras como marco complementario"):
    add_paragraph(p_text)

# ── 2.2 Estado de la cuestión ──
add_heading("Estado de la cuestión", level=2)

# Agentes LLM (§3.3)
add_paragraph(
    "2.2.1. Del asesoramiento automatizado a los agentes de inversión "
    "basados en modelos de lenguaje de gran escala"
)
for p_text in sec("3.3. Del asesoramiento automatizado a los agentes de "
                   "inversión basados en LLM"):
    add_paragraph(p_text)

# Síntesis (§3.4)
add_blank()
add_paragraph(
    "2.2.2. Síntesis del estado de la cuestión sobre sistemas existentes"
)
for p_text in sec("3.4. Síntesis del estado de la cuestión sobre lo que ya existe"):
    add_paragraph(p_text)

# Tabla comparativa 1 (si existe)
if len(tablas_orig) > 0:
    add_blank()
    add_paragraph(
        "Tabla 1. Comparativa de las principales familias de soluciones "
        "existentes y su relación con el sistema propuesto."
    )
    t0 = tablas_orig[0]
    if len(t0) > 1:
        add_table(t0[0], t0[1:])
    add_blank()

# NLP financiero — Del sentiment analysis al modelado semántico (§4.1)
add_blank()
add_paragraph(
    "2.2.3. Del análisis de sentimiento al modelado semántico del evento"
)
for p_text in sec("4.1. Del sentiment analysis al modelado semántico del evento"):
    add_paragraph(p_text)

# Datasets, benchmarks y extracción de eventos (§4.2)
add_blank()
add_paragraph(
    "2.2.4. Conjuntos de datos, bancos de pruebas y extracción de "
    "eventos financieros"
)
for p_text in sec("4.2. Datasets, benchmarks y extracción de eventos financieros"):
    add_paragraph(p_text)

# Relevancia contextual (§4.3)
add_blank()
add_paragraph("2.2.5. Relevancia contextual: por qué la cartera cambia el problema")
for p_text in sec("4.3. Relevancia contextual: por qué la cartera cambia el problema"):
    add_paragraph(p_text)

# Noticias y reacción del mercado (§4.4)
add_blank()
add_paragraph(
    "2.2.6. Noticias, reacción del mercado y límites de inferencia"
)
for p_text in sec("4.4. Noticias, reacción del mercado y límites de inferencia"):
    add_paragraph(p_text)

# Análisis comparativo (§6)
add_blank()
add_paragraph("2.2.7. Análisis comparativo de trabajos previos")
for p_text in sec("6. Análisis comparativo de trabajos previos"):
    add_paragraph(p_text)

# Criterios de comparación (§6.1)
for p_text in sec("6.1. Criterios de comparación"):
    add_paragraph(p_text)

# Tabla comparativa 2 (si existe)
if len(tablas_orig) > 1:
    add_blank()
    add_paragraph(
        "Tabla 2. Comparativa de trabajos previos según cinco criterios."
    )
    t1 = tablas_orig[1]
    if len(t1) > 1:
        add_table(t1[0], t1[1:])
    add_blank()

# ── 2.3 Conclusiones del estado de la cuestión ──
add_heading("Conclusiones", level=2)

# Implicaciones para el TFM (§4.5)
for p_text in sec("4.5. Implicaciones para este TFM"):
    add_paragraph(p_text)

# Hueco de investigación (§7) — completo aquí como cierre
add_blank()
add_paragraph(
    "La identificación de este hueco de investigación se apoya en cinco "
    "observaciones convergentes:"
)
hueco = sec("7. Hueco de investigación detectado")
for p_text in hueco:
    add_paragraph(p_text)

add_paragraph(
    "Estas conclusiones configuran el hueco de investigación al que "
    "responde el presente trabajo y fundamentan los objetivos que se "
    "formulan en el capítulo siguiente."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  CAPÍTULO 3: OBJETIVOS CONCRETOS Y METODOLOGÍA DE TRABAJO
# ══════════════════════════════════════════════════════════════════

add_heading("Objetivos concretos y metodología de trabajo", level=1)

add_paragraph(
    "Este capítulo formula los objetivos del trabajo y describe la "
    "metodología seguida para alcanzarlos. Los objetivos se plantean "
    "siguiendo el criterio SMART (Doran, 1981): específicos, medibles, "
    "alcanzables, relevantes y temporalmente acotados."
)

# ── 3.1 Objetivo general ──
add_heading("Objetivo general", level=2)

add_paragraph(
    "Diseñar, implementar y evaluar un sistema de alertas inteligentes "
    "por noticias financieras que, dado un flujo heterogéneo de fuentes "
    "textuales y una cartera de inversión declarada por el usuario, "
    "detecte eventos materialmente relevantes, estime la dirección y "
    "severidad de su impacto, y genere alertas explicables con control "
    "de ruido y trazabilidad de fuente, alcanzando un F1 superior a 0,80 "
    "en la detección de relevancia sobre un corpus de evaluación "
    "etiquetado manualmente."
)

# ── 3.2 Objetivos específicos ──
add_heading("Objetivos específicos", level=2)

add_paragraph(
    "El objetivo general se descompone en los siguientes objetivos "
    "específicos:"
)

objetivos = [
    (
        "OE1. Revisar y sintetizar el estado de la cuestión sobre "
        "robo-advisors, procesamiento de lenguaje natural financiero, "
        "extracción de eventos y agentes basados en modelos de lenguaje "
        "de gran escala, identificando el hueco de investigación que "
        "justifica la integración propuesta."
    ),
    (
        "OE2. Diseñar e implementar una cadena de procesamiento modular "
        "de ingesta, preprocesado y enriquecimiento de noticias "
        "financieras procedentes de fuentes heterogéneas (RSS, "
        "interfaz de programación de aplicaciones (API, por sus siglas "
        "en inglés) de la SEC EDGAR, CNMV, NewsAPI), con detección de "
        "idioma, traducción automática y reconocimiento de entidades "
        "nombradas (NER, por sus siglas en inglés)."
    ),
    (
        "OE3. Desarrollar un módulo de relevancia híbrido que combine "
        "reglas explícitas con representaciones vectoriales densas "
        "(embeddings) semánticas para evaluar la pertinencia de cada "
        "noticia en función de la composición concreta de la cartera "
        "del usuario."
    ),
    (
        "OE4. Implementar la clasificación automática de eventos "
        "financieros y la estimación de impacto mediante análisis de "
        "sentimiento con FinBERT, inferencia de lenguaje natural (NLI) "
        "y análisis contextual opcional con modelos de lenguaje de gran "
        "escala."
    ),
    (
        "OE5. Construir un conjunto de datos etiquetado manualmente con "
        "al menos 40 noticias, tres carteras de referencia y etiquetas "
        "de relevancia, tipo de evento, dirección, severidad y activos "
        "afectados, que sirva como banco de pruebas reproducible para "
        "la evaluación del sistema."
    ),
    (
        "OE6. Evaluar cuantitativamente el sistema mediante un estudio "
        "de ablación que compare al menos cuatro variantes de la cadena "
        "de procesamiento, midiendo precisión, exhaustividad, F1 y "
        "error absoluto medio ordinal en cada tarea."
    ),
]
for obj in objetivos:
    add_list_item(obj)

# ── 3.3 Metodología del trabajo ──
add_heading("Metodología del trabajo", level=2)

# Del original §9 (Metodología prevista)
for p_text in sec("9. Metodología prevista"):
    add_paragraph(p_text)

# Fases detalladas
add_paragraph(
    "A continuación se describen las seis fases de desarrollo del sistema:"
)

fases = [
    ("9.1. Fase 1: alcance experimental, activos y taxonomía de eventos",
     "Fase 1: Alcance experimental, activos y taxonomía de eventos"),
    ("9.2. Fase 2: adquisición de datos y trazabilidad",
     "Fase 2: Adquisición de datos y trazabilidad"),
    ("9.3. Fase 3: preprocesado NLP y enriquecimiento",
     "Fase 3: Preprocesado y enriquecimiento"),
    ("9.4. Fase 4: relevancia por cartera",
     "Fase 4: Relevancia por cartera"),
    ("9.5. Fase 5: clasificación de eventos e impacto",
     "Fase 5: Clasificación de eventos e impacto"),
    ("9.6. Fase 6: motor de alertas y evaluación",
     "Fase 6: Motor de alertas y evaluación"),
]

for sec_name, titulo_fase in fases:
    add_blank()
    add_paragraph(titulo_fase)
    for p_text in sec(sec_name):
        add_paragraph(p_text)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  CAPÍTULO 4: DESARROLLO ESPECÍFICO DE LA CONTRIBUCIÓN
#  (Tipo 2: Desarrollo de software)
# ══════════════════════════════════════════════════════════════════

add_heading("Desarrollo específico de la contribución", level=1)

add_paragraph(
    "Este capítulo describe en detalle la contribución técnica del trabajo: "
    "la identificación de requisitos del sistema, la arquitectura y módulos "
    "de la herramienta desarrollada, y la evaluación cuantitativa de su "
    "rendimiento."
)

# ── 4.1 Identificación de requisitos ──
add_heading("Identificación de requisitos", level=2)

# De §8 (Propuesta del sistema) + §8.1 (Salida esperada)
add_paragraph(
    "Los requisitos del sistema se derivan directamente del hueco de "
    "investigación identificado en el capítulo 2 y de los objetivos "
    "formulados en el capítulo 3. Se distinguen requisitos funcionales "
    "y no funcionales."
)

add_paragraph("Requisitos funcionales:")
reqs_func = [
    "RF1. El sistema debe permitir al usuario declarar una cartera de "
    "inversión con tickers, sectores y geografías.",
    "RF2. Debe adquirir noticias de forma automática desde fuentes "
    "heterogéneas (RSS, SEC EDGAR, CNMV, NewsAPI, Alpha Vantage).",
    "RF3. Debe preprocesar las noticias: limpieza, detección de idioma, "
    "traducción automática y reconocimiento de entidades.",
    "RF4. Debe evaluar la relevancia de cada noticia respecto a la "
    "cartera del usuario mediante reglas explícitas y similitud semántica.",
    "RF5. Debe clasificar el tipo de evento financiero y estimar la "
    "dirección y severidad del impacto.",
    "RF6. Debe generar alertas explicables con trazabilidad de fuente, "
    "deduplicación semántica y control de ruido.",
    "RF7. Debe ofrecer opcionalmente un análisis contextual enriquecido "
    "mediante modelos de lenguaje de gran escala.",
    "RF8. Debe proporcionar un módulo de asesoramiento de inversiones "
    "con perfilado de riesgo MiFID.",
    "RF9. Debe ofrecer métricas de cartera (Sharpe, Sortino, VaR, "
    "drawdown) y datos de mercado en tiempo real.",
]
for req in reqs_func:
    add_list_item(req)

add_blank()
add_paragraph("Requisitos no funcionales:")
reqs_nofunc = [
    "RNF1. Arquitectura modular que permita la sustitución independiente "
    "de componentes.",
    "RNF2. Funcionamiento sin dependencia obligatoria de servicios de "
    "pago (degradación elegante sin claves de API).",
    "RNF3. Soporte multilingüe (español e inglés como mínimo).",
    "RNF4. Interfaz web responsiva para la gestión de carteras y "
    "visualización de alertas.",
    "RNF5. Despliegue reproducible mediante contenedores Docker.",
]
for req in reqs_nofunc:
    add_list_item(req)

# Salida esperada de la alerta (§8.1)
add_blank()
add_paragraph("Salida esperada de una alerta:")
for p_text in sec("8.1. Salida esperada de la alerta"):
    add_paragraph(p_text)

# Ejemplo de alerta (tabla 3 si existe)
if len(tablas_orig) > 2:
    add_blank()
    t2 = tablas_orig[2]
    for row in t2:
        for cell in row:
            if cell.strip():
                add_paragraph(cell.strip())

# ── 4.2 Descripción de la herramienta desarrollada ──
add_heading("Descripción de la herramienta software desarrollada", level=2)

add_paragraph(
    "El sistema se ha implementado siguiendo una arquitectura modular "
    "organizada en una cadena de procesamiento de siete etapas. La "
    "Figura 1 muestra el diagrama general de la arquitectura."
)

add_paragraph("[Figura 1. Arquitectura general del sistema InvestAIlert. "
              "Fuente: elaboración propia.]")
add_blank()

# Stack tecnológico
add_paragraph("Tecnologías empleadas:")
add_paragraph(
    "El sistema se ha desarrollado en Python 3.12 con FastAPI como marco "
    "de la interfaz de programación de aplicaciones REST asíncrona y "
    "Motor como controlador asíncrono de MongoDB. La interfaz web se ha "
    "construido con Next.js 16.2 y React 19, empleando Tailwind CSS para "
    "el diseño responsivo."
)

add_paragraph("Modelos de procesamiento de lenguaje natural:")
modelos_nlp = [
    "Análisis de sentimiento financiero: FinBERT (ProsusAI/finbert), "
    "modelo preentrenado sobre textos financieros basado en la "
    "arquitectura BERT.",
    "Representaciones vectoriales densas: Sentence-Transformers "
    "(all-MiniLM-L6-v2), para similitud semántica y deduplicación. El "
    "término embeddings se emplea de forma generalizada en la comunidad "
    "investigadora al no existir un equivalente consolidado en español "
    "(Mikolov et al., 2013; Reimers y Gurevych, 2019).",
    "Reconocimiento de entidades: spaCy (en_core_web_sm), con "
    "traducción automática de textos en español mediante deep-translator.",
    "Clasificación de eventos: BART-MNLI (facebook/bart-large-mnli), "
    "modelo de inferencia de lenguaje natural en configuración de "
    "clasificación sin ejemplos previos (zero-shot). Reemplaza la "
    "dependencia de modelos de lenguaje de gran escala para esta tarea, "
    "funcionando de forma local y sin coste.",
    "Análisis contextual: multiproveedor (OpenAI, GitHub Models, "
    "HuggingFace, Ollama) mediante interfaz unificada AsyncOpenAI. "
    "Funcionalidad opcional con degradación elegante.",
]
for m in modelos_nlp:
    add_list_item(m)

add_blank()
add_paragraph("Módulos del sistema:")
add_paragraph(
    "A continuación se describe brevemente cada módulo del sistema. "
    "El código fuente completo está disponible en el repositorio "
    "indicado en el Anexo A."
)

modulos = [
    ("Módulo de cartera (portfolio)", "Gestión completa de carteras de "
     "inversión: creación, lectura, actualización y eliminación. "
     "Normalización de tickers, sectores y geografías."),
    ("Módulo de ingesta (ingestion)", "Adquisición automática de noticias "
     "desde cinco fuentes: RSS (feeds financieros internacionales), "
     "CNMV (hechos relevantes del regulador español), NewsAPI, "
     "Alpha Vantage y SEC EDGAR (filings de la Comisión de Bolsa y "
     "Valores de Estados Unidos)."),
    ("Módulo de preprocesado (nlp)", "Limpieza textual, detección de "
     "idioma, traducción automática español→inglés, análisis de "
     "sentimiento con FinBERT y reconocimiento de entidades con spaCy."),
    ("Módulo de relevancia (relevance)", "Evaluación híbrida de la "
     "pertinencia de cada noticia: capa de reglas explícitas (menciones "
     "directas a tickers y nombres de empresas) y capa semántica "
     "(similitud coseno entre representaciones vectoriales de la "
     "noticia y los activos de la cartera)."),
    ("Módulo de eventos (events)", "Clasificación automática del tipo "
     "de evento financiero entre 12 categorías predefinidas mediante "
     "inferencia de lenguaje natural sin ejemplos previos con BART-MNLI."),
    ("Módulo de impacto (impact)", "Estimación determinista de la "
     "dirección (alcista/bajista/neutral) y la severidad (1-5) del "
     "impacto, con posibilidad de enriquecimiento contextual mediante "
     "modelos de lenguaje de gran escala."),
    ("Motor de alertas (alerts)", "Generación de alertas con puntuación "
     "compuesta, deduplicación semántica mediante representaciones "
     "vectoriales con persistencia en MongoDB, y control antirruido."),
    ("Módulo de notificaciones (notifications)", "Envío de alertas por "
     "correo electrónico (SMTP con plantilla HTML) y webhook HTTP "
     "(compatible con Slack, Discord y Telegram)."),
    ("Módulo de asesoramiento (advisor)", "Cuestionario MiFID de 10 "
     "preguntas para perfilado de riesgo, análisis de diversificación "
     "(índice Herfindahl-Hirschman) y generación de informes con "
     "modelos de lenguaje de gran escala o alternativa determinista."),
    ("Módulo de mercado (market)", "Datos de mercado en tiempo real "
     "mediante yfinance: búsqueda de activos, precios actuales e "
     "históricos OHLCV."),
    ("Módulo de analítica (analytics)", "Métricas de rendimiento de "
     "cartera mediante quantstats: ratio de Sharpe, ratio de Sortino, "
     "valor en riesgo (VaR), máxima caída (drawdown), alfa y beta."),
    ("Programador (scheduler)", "Orquestación de tareas periódicas con "
     "APScheduler: ingesta cada 15 minutos, generación de alertas cada "
     "20 minutos, limpieza diaria de registros antiguos."),
]
for nombre, desc in modulos:
    add_blank()
    add_paragraph(f"{nombre}. {desc}")

add_blank()
add_paragraph("Interfaz web:")
add_paragraph(
    "La interfaz web proporciona seis vistas principales: panel de "
    "control con indicadores clave de rendimiento financiero, gestión "
    "de carteras con autocompletado de tickers, visualización de "
    "noticias procesadas, motor de alertas con filtros, módulo de "
    "asesoramiento con cuestionario interactivo y configuración del "
    "sistema."
)
add_paragraph("[Figura 2. Captura de pantalla del panel de control. "
              "Fuente: elaboración propia.]")

# ── 4.3 Evaluación ──
add_heading("Evaluación", level=2)

add_paragraph(
    "La evaluación del sistema se ha abordado desde una perspectiva "
    "cuantitativa, mediante un estudio de ablación que compara cuatro "
    "variantes de la cadena de procesamiento sobre un corpus de "
    "evaluación etiquetado manualmente."
)

# Dataset como contribución
add_blank()
add_paragraph("4.3.1. Conjunto de datos etiquetado como contribución")
add_paragraph(
    "Una de las contribuciones del presente trabajo es la construcción "
    "de un conjunto de datos etiquetado específicamente diseñado para "
    "evaluar sistemas de alertas financieras personalizadas. Este "
    "recurso, que no existía previamente en la literatura consultada, "
    "se compone de los siguientes elementos:"
)
dataset_items = [
    "40 noticias financieras reales, redactadas en inglés y español.",
    "3 carteras de referencia con composiciones diversas: tecnológica "
    "estadounidense (AAPL, MSFT, NVDA, TSLA), ibérica diversificada "
    "(SAN.MC, BBVA.MC, ITX.MC, IBE.MC) y energética global "
    "(XOM, SHEL, REP.MC).",
    "Etiquetas manuales para cada par noticia-cartera: relevancia "
    "(binaria), activos afectados (conjunto), tipo de evento (12 "
    "categorías), dirección (alcista/bajista/neutral) y severidad "
    "(escala ordinal 1-5).",
    "5 noticias negativas (irrelevantes para todas las carteras) "
    "para evaluar la tasa de falsos positivos.",
    "Casos difíciles de relevancia indirecta: regulación sectorial, "
    "competidores, proveedores.",
    "Cobertura de los 12 tipos de evento definidos en la taxonomía "
    "del sistema.",
]
for item in dataset_items:
    add_list_item(item)

add_paragraph(
    "El conjunto de datos se distribuye en formato JSONL con esquema "
    "validado mediante modelos Pydantic y está disponible en el "
    "repositorio del proyecto (véase Anexo A)."
)

# Marco de evaluación
add_blank()
add_paragraph("4.3.2. Marco de evaluación y métricas")
add_paragraph(
    "La evaluación emplea un conjunto de métricas complementarias "
    "adaptadas a la naturaleza de cada tarea:"
)
metricas = [
    "Relevancia (clasificación binaria): precisión, exhaustividad "
    "(recall), F1 y exactitud (accuracy).",
    "Activos afectados (predicción de conjuntos): índice de Jaccard, "
    "precisión micro, exhaustividad micro y F1 micro.",
    "Tipo de evento (clasificación multiclase): precisión, "
    "exhaustividad y F1 por clase; promedios macro y ponderado.",
    "Severidad (escala ordinal 1-5): error absoluto medio (MAE) "
    "sobre la escala de cinco niveles.",
]
for m in metricas:
    add_list_item(m)

add_paragraph(
    "Todas las métricas se han implementado en Python puro, sin "
    "dependencias externas como scikit-learn, para garantizar la "
    "reproducibilidad y la transparencia del cálculo."
)

# Estudio de ablación
add_blank()
add_paragraph("4.3.3. Estudio de ablación")
add_paragraph(
    "El estudio de ablación compara cuatro variantes del sistema, "
    "construidas mediante la activación progresiva de módulos:"
)
variantes = [
    "Variante «reglas» (baseline): solo coincidencia exacta de "
    "tickers y nombres de empresas.",
    "Variante «híbrida»: reglas + similitud semántica mediante "
    "representaciones vectoriales densas.",
    "Variante «híbrida con NLI»: reglas + semántica + clasificación "
    "de eventos con inferencia de lenguaje natural.",
    "Variante «completa»: todos los módulos anteriores + análisis "
    "contextual con modelo de lenguaje de gran escala.",
]
for v in variantes:
    add_list_item(v)

# Resultados
add_blank()
add_paragraph("4.3.4. Resultados preliminares")
add_paragraph(
    "La Tabla 3 presenta los resultados de la variante basada en "
    "reglas, que constituye el baseline del sistema."
)

add_paragraph(
    "Tabla 3. Resultados de la variante «reglas» sobre el corpus "
    "de evaluación."
)
add_table(
    ["Tarea", "Métrica", "Valor"],
    [
        ["Relevancia", "Precisión", "1,000"],
        ["Relevancia", "Exhaustividad", "0,879"],
        ["Relevancia", "F1", "0,935"],
        ["Relevancia", "Exactitud", "0,925"],
        ["Activos afectados", "F1 micro", "1,000"],
        ["Activos afectados", "Jaccard", "1,000"],
        ["Severidad", "MAE ordinal", "[pendiente]"],
    ],
)

add_blank()
add_paragraph(
    "[NOTA: Los resultados de las variantes híbrida, híbrida con NLI "
    "y completa se incorporarán a medida que se ejecuten las pruebas "
    "correspondientes. El estudio de ablación completo requiere "
    "aproximadamente 10-15 minutos de procesamiento en CPU.]"
)

add_paragraph(
    "Estos resultados preliminares son alentadores: la variante más "
    "sencilla ya supera el umbral de F1=0,80 establecido en el "
    "objetivo general, con una precisión perfecta que indica ausencia "
    "total de falsos positivos. La exhaustividad de 0,879 señala que "
    "el sistema pierde un 12,1% de noticias relevantes, "
    "correspondientes a casos de relevancia indirecta que requieren "
    "las capas semánticas o de modelos de lenguaje de gran escala para "
    "ser detectados."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  CAPÍTULO 5: CONCLUSIONES Y TRABAJO FUTURO
# ══════════════════════════════════════════════════════════════════

add_heading("Conclusiones y trabajo futuro", level=1)

# ── 5.1 Conclusiones ──
add_heading("Conclusiones", level=2)

# Del original §10
for p_text in sec("10. Conclusiones y plan de trabajo"):
    add_paragraph(p_text)

add_paragraph(
    "Adicionalmente, la construcción de un conjunto de datos etiquetado "
    "bilingüe para la evaluación de sistemas de alertas financieras "
    "personalizadas constituye una contribución independiente que puede "
    "ser reutilizada por otros investigadores del área."
)

# ── 5.2 Líneas de trabajo futuro ──
add_heading("Líneas de trabajo futuro", level=2)

lineas_futuro = [
    "Ampliación del conjunto de datos de evaluación con más noticias, "
    "más carteras y anotadores múltiples para medir la concordancia "
    "inter-anotador (kappa de Cohen).",
    "Ejecución del estudio de ablación completo con las cuatro "
    "variantes y análisis estadístico de las diferencias.",
    "Incorporación de un modelo de reconocimiento de entidades "
    "financiero especializado (por ejemplo, basado en SENTiVENT o "
    "entrenado con datos específicos del dominio).",
    "Integración de un estudio de evento simplificado que mida el "
    "retorno anormal acumulado (CAR) en ventanas temporales alrededor "
    "de las alertas generadas.",
    "Caché de representaciones vectoriales de cartera para evitar "
    "recálculos y mejorar la latencia del sistema.",
    "Implementación de autenticación con JWT y autorización por "
    "usuario para soportar múltiples inversores simultáneos.",
    "Exploración de modelos multilingües nativos (paraphrase-"
    "multilingual-MiniLM-L12-v2) para eliminar la dependencia de "
    "traducción automática.",
    "Calibración de las probabilidades de confianza del modelo "
    "mediante escalado de Platt o regresión isotónica.",
    "Pruebas de usabilidad con usuarios expertos para evaluar la "
    "aplicabilidad práctica del sistema.",
]

for linea in lineas_futuro:
    add_list_item(linea)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  REFERENCIAS BIBLIOGRÁFICAS
# ══════════════════════════════════════════════════════════════════

add_heading("Referencias bibliográficas", level=1)

# Del original + nuevas referencias
referencias = sec("Referencias bibliográficas")

# Añadir las referencias nuevas (alucinaciones LLM, embeddings)
nuevas_refs = [
    "Doran, G. T. (1981). There's a S.M.A.R.T. way to write "
    "management's goals and objectives. Management Review (AMA FORUM), "
    "70, 35-36.",
    "Ji, Z., Lee, N., Frieske, R., Yu, T., Su, D., Xu, Y., Ishii, E., "
    "Bang, Y., Madotto, A., y Fung, P. (2023). Survey of Hallucination "
    "in Natural Language Generation. ACM Computing Surveys, 55(12), "
    "1–38. https://doi.org/10.1145/3571730",
    "Mikolov, T., Sutskever, I., Chen, K., Corrado, G. S., y Dean, J. "
    "(2013). Distributed Representations of Words and Phrases and their "
    "Compositionality. Advances in Neural Information Processing "
    "Systems, 26.",
    "Reimers, N., y Gurevych, I. (2019). Sentence-BERT: Sentence "
    "Embeddings using Siamese BERT-Networks. Proceedings of the 2019 "
    "Conference on Empirical Methods in Natural Language Processing "
    "(EMNLP). https://doi.org/10.18653/v1/D19-1410",
]

# Combinar y ordenar alfabéticamente
todas_refs = list(set(referencias + nuevas_refs))
todas_refs.sort(key=lambda x: x.lower())

for ref in todas_refs:
    if ref.strip():
        add_paragraph(ref)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
#  ANEXO A: CÓDIGO FUENTE Y DATOS ANALIZADOS
# ══════════════════════════════════════════════════════════════════

add_heading("Anexo A. Código fuente y datos analizados", level=1)

add_paragraph(
    "El código fuente completo del sistema InvestAIlert se encuentra "
    "alojado en el siguiente repositorio, del que el autor es único "
    "propietario y contribuidor:"
)
add_paragraph("[URL del repositorio de GitHub]")
add_blank()
add_paragraph(
    "El repositorio contiene la totalidad del código desarrollado "
    "durante el Trabajo Fin de Máster, incluyendo:"
)
anexo_items = [
    "Backend completo (Python 3.12, FastAPI): main.py, config.py y "
    "módulos en modules/.",
    "Frontend (Next.js 16.2, React 19): en frontend/src/.",
    "Módulo de evaluación: en evaluation/, incluyendo el conjunto de "
    "datos etiquetado (dataset.jsonl), las carteras de referencia "
    "(portfolios.json), las métricas (metrics.py) y el ejecutor del "
    "estudio de ablación (run_ablation.py).",
    "Pruebas automatizadas: 137 pruebas unitarias en tests/.",
    "Configuración de despliegue: Dockerfile, docker-compose.yml.",
    "Documentación técnica: README.md, NOTAS_TECNICAS.md, "
    "HERRAMIENTAS_OPEN_SOURCE.md.",
]
for item in anexo_items:
    add_list_item(item)

add_blank()
add_paragraph(
    "Los datos utilizados para la evaluación (conjunto de datos "
    "etiquetado de 40 noticias) también se encuentran disponibles "
    "en el directorio evaluation/ del repositorio."
)


# ──────────────────────────────────────────────────────────────────
# 7. Guardar el documento
# ──────────────────────────────────────────────────────────────────

doc.save(SALIDA)
print(f"\n✅ Documento generado: {SALIDA}")
print(f"   Secciones del cuerpo: 5 capítulos + resumen + abstract + anexo")
print(f"   Directrices aplicadas:")
print(f"     - Anglicismos sustituidos por equivalentes en español")
print(f"     - Acrónimos definidos en primera mención")
print(f"     - Embeddings justificado como anglicismo necesario")
print(f"     - Párrafo sobre alucinaciones de LLM en Motivación")
print(f"     - Conjunto de datos de evaluación como contribución")
print(f"     - Estructura según plantilla_individual (Tipo 2: Desarrollo SW)")
