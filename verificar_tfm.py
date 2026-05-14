"""Verificación exhaustiva del TFM contra plantilla + directrices del tutor."""
from docx import Document
import re

doc = Document(r"TFM_alertas_inversion_reestructurado.docx")
full_text = " ".join(p.text for p in doc.paragraphs)

# Extraer estructura de headings
headings = []
for p in doc.paragraphs:
    style = p.style.name if p.style else ""
    if "Heading" in style or "Title" in style:
        headings.append((style, p.text.strip()))

print("=" * 60)
print(" VERIFICACIÓN: ESTRUCTURA vs PLANTILLA_INDIVIDUAL")
print("=" * 60)
print()

estructura_requerida = [
    "Resumen",
    "Abstract",
    "Introducción",
    "Motivación",
    "Planteamiento del trabajo",
    "Estructura del trabajo",
    "Contexto y estado",
    "Contexto del problema",
    "Estado de la cuestión",
    "Conclusiones",
    "Objetivos concretos y metodología",
    "Objetivo general",
    "Objetivos específicos",
    "Metodología del trabajo",
    "Desarrollo específico de la contribución",
    "Conclusiones y trabajo futuro",
    "Líneas de trabajo futuro",
    "Referencias bibliográficas",
    "Anexo",
]

for req in estructura_requerida:
    found = any(req.lower() in h[1].lower() for h in headings)
    print(f'  {"PASS" if found else "FAIL"} {req}')

print()
print("=" * 60)
print(" VERIFICACIÓN: DIRECTRICES DEL TUTOR")
print("=" * 60)

# 1. Anglicismos
print()
print("1. ANGLICISMOS MINIMIZADOS:")
bad_terms = ["sentiment analysis", "portfolio optimization", "pipeline", "dataset", "framework", "feedback"]
for term in bad_terms:
    count = len(re.findall(r"\b" + re.escape(term) + r"\b", full_text, re.I))
    print(f'   {"PASS" if count == 0 else "FAIL"} "{term}": {count} (debe ser 0)')

good_terms = {
    "análisis de sentimiento": 0,
    "cadena de procesamiento": 0,
    "conjunto de datos": 0,
    "bancos de pruebas": 0,
}
print("   Equivalentes en español:")
for term in good_terms:
    count = len(re.findall(term, full_text, re.I))
    print(f'   {"PASS" if count > 0 else "FAIL"} "{term}": {count}')

# 2. Acrónimos
print()
print("2. ACRONIMOS DEFINIDOS EN PRIMERA MENCION:")
acronimos = {
    "NLP": "procesamiento de lenguaje natural (NLP, por sus siglas en inglés)",
    "LLM": "modelo de lenguaje de gran escala (LLM, por sus siglas en inglés)",
    "NER": "reconocimiento de entidades nombradas (NER, por sus siglas en inglés)",
    "NLI": "inferencia de lenguaje natural (NLI, por sus siglas en inglés)",
    "API": "interfaz de programación de aplicaciones (API, por sus siglas en inglés)",
    "TFM": "Trabajo Fin de Máster (TFM)",
}
for sigla, definicion in acronimos.items():
    found = definicion in full_text
    print(f'   {"PASS" if found else "FAIL"} {sigla}')

# 3. Estado de la cuestión
print()
print('3. "ESTADO DE LA CUESTION" (no "estado del arte"):')
edo_cuestion = len(re.findall(r"estado de la cuestión", full_text, re.I))
edo_arte = len(re.findall(r"estado del arte", full_text, re.I))
print(f'   {"PASS" if edo_cuestion > 0 else "FAIL"} "estado de la cuestion": {edo_cuestion}')
print(f'   {"PASS" if edo_arte == 0 else "WARN"} "estado del arte": {edo_arte} (debe ser 0)')

# 4. Alucinaciones LLM en motivación
print()
print("4. ALUCINACIONES LLM EN MOTIVACION:")
motivacion_start = full_text.find("Motivación")
planteamiento_start = full_text.find("Planteamiento del trabajo")
if motivacion_start > 0 and planteamiento_start > motivacion_start:
    sec_text = full_text[motivacion_start:planteamiento_start]
    a = "alucinaciones" in sec_text
    b = "no sustentada en evidencia" in sec_text
    print(f'   {"PASS" if a else "FAIL"} Menciona alucinaciones')
    print(f'   {"PASS" if b else "FAIL"} Info no basada en evidencias')

# 5. Embeddings justificado
print()
print("5. EMBEDDINGS JUSTIFICADO:")
j = "no existir un equivalente consolidado en español" in full_text
c = "Mikolov" in full_text and "Reimers" in full_text
print(f'   {"PASS" if j else "FAIL"} Justificacion del anglicismo')
print(f'   {"PASS" if c else "FAIL"} Citas (Mikolov, Reimers)')

# 6. Dataset como contribución
print()
print("6. DATASET COMO CONTRIBUCION:")
contrib = "contribuci" in full_text.lower() and "conjunto de datos" in full_text.lower()
corpus = "40 noticias" in full_text
biling = "bilingüe" in full_text
print(f'   {"PASS" if contrib else "FAIL"} Presentado como contribucion')
print(f'   {"PASS" if corpus else "FAIL"} Corpus 40 noticias')
print(f'   {"PASS" if biling else "FAIL"} Bilingue')

# 7. Tipo 2 - Desarrollo de software
print()
print("7. TIPO 2 DESARROLLO DE SOFTWARE:")
r = any("requisitos" in h[1].lower() for h in headings)
s = any("herramienta" in h[1].lower() or "software" in h[1].lower() for h in headings)
e = any("evaluación" in h[1].lower() for h in headings)
print(f'   {"PASS" if r else "FAIL"} Identificacion de requisitos')
print(f'   {"PASS" if s else "FAIL"} Descripcion de la herramienta')
print(f'   {"PASS" if e else "FAIL"} Evaluacion')

# 8. Objetivos SMART
print()
print("8. OBJETIVOS SMART:")
sm = "SMART" in full_text
dr = "Doran" in full_text
oe = full_text.count("OE")
print(f'   {"PASS" if sm else "FAIL"} Metodologia SMART mencionada')
print(f'   {"PASS" if dr else "FAIL"} Cita Doran (1981)')
print(f'   {"PASS" if oe >= 6 else "FAIL"} Objetivos especificos: {oe}')

print()
print("=" * 60)
print(" VERIFICACION COMPLETA")
print("=" * 60)
